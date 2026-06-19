import logging
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from pypdf import PdfReader

from ..config import get_settings

try:
    import pypdfium2 as pdfium
except Exception:  # pragma: no cover - optional dependency
    pdfium = None

try:
    from PIL import Image
except Exception:  # pragma: no cover - optional dependency
    Image = None

try:
    import pytesseract
    from pytesseract import TesseractError
except Exception:  # pragma: no cover - optional dependency
    pytesseract = None


settings = get_settings()
logger = logging.getLogger(__name__)


class OCRService:
    @staticmethod
    def _resolve_languages(source_language: Optional[str]) -> str:
        normalized = (source_language or '').strip().lower()
        language_map = {
            'hindi': 'hin+eng',
            'hi': 'hin+eng',
            'english': 'eng',
            'en': 'eng',
            'marathi': 'mar+eng',
            'mr': 'mar+eng',
        }
        if normalized in language_map:
            return language_map[normalized]

        configured = (settings.ocr_languages or '').strip()
        return configured or 'eng'

    @staticmethod
    def _configure_tesseract() -> bool:
        if pytesseract is None:
            return False

        if settings.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

        try:
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False

    @staticmethod
    def _ocr_image(image: "Image.Image", languages: str) -> str:
        if pytesseract is None:
            return ''

        ocr_languages = languages.strip() or 'eng'
        tesseract_config = (settings.ocr_tesseract_config or '').strip()

        # Improve OCR robustness for scanned pages by normalizing to grayscale and
        # upscaling small images before text recognition.
        normalized = image.convert('L')
        width, height = normalized.size
        if max(width, height) < 1800:
            normalized = normalized.resize((width * 2, height * 2))

        try:
            return (
                pytesseract.image_to_string(
                    normalized,
                    lang=ocr_languages,
                    config=tesseract_config,
                )
                or ''
            ).strip()
        except TesseractError as exc:
            if ocr_languages != 'eng':
                logger.warning(
                    'Tesseract OCR failed for languages "%s" (%s). Retrying with "eng".',
                    ocr_languages,
                    exc,
                )
                try:
                    return (
                        pytesseract.image_to_string(
                            normalized,
                            lang='eng',
                            config=tesseract_config,
                        )
                        or ''
                    ).strip()
                except Exception:
                    return ''
            return ''
        except Exception:
            return ''

    @staticmethod
    def _ocr_pdf_pages(file_path: str, languages: str) -> str:
        if pdfium is None or Image is None:
            return ''

        if not OCRService._configure_tesseract():
            return ''

        try:
            pdf = pdfium.PdfDocument(file_path)
        except Exception:
            return ''

        chunks: list[str] = []
        page_count = min(len(pdf), max(settings.ocr_max_pages, 1))

        try:
            for page_index in range(page_count):
                page = pdf[page_index]
                render = page.render(scale=max(settings.ocr_render_scale, 1.0))
                pil_image = render.to_pil()
                text = OCRService._ocr_image(pil_image, languages)
                if text:
                    chunks.append(text)
        except Exception:
            return ''
        finally:
            try:
                pdf.close()
            except Exception:
                pass

        return '\n\n'.join(chunks).strip()

    @staticmethod
    def _ocr_standalone_image(file_path: str, languages: str) -> str:
        if Image is None:
            return ''

        if not OCRService._configure_tesseract():
            return ''

        try:
            with Image.open(file_path) as image:
                return OCRService._ocr_image(image, languages)
        except Exception:
            return ''

    @staticmethod
    def extract_text(file_path: str, source_language: Optional[str] = None) -> str:
        path = Path(file_path)
        suffix = path.suffix.lower()
        languages = OCRService._resolve_languages(source_language)

        if suffix in {'.txt'}:
            return path.read_text(encoding='utf-8', errors='ignore')

        if suffix in {'.docx'}:
            document = DocxDocument(file_path)
            text = '\n'.join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
            return text or 'No readable text found in DOCX.'

        if suffix in {'.pdf'}:
            reader = PdfReader(file_path)
            chunks: list[str] = []
            for page in reader.pages:
                page_text = page.extract_text() or ''
                if page_text.strip():
                    chunks.append(page_text.strip())
            if chunks:
                return '\n\n'.join(chunks)

            ocr_text = OCRService._ocr_pdf_pages(file_path, languages)
            if ocr_text:
                return ocr_text

            return (
                'No readable text found in PDF. If this is a scanned/image-based PDF, '
                'install and configure OCR dependencies (Tesseract with Hindi language data).'
            )

        if suffix in {'.png', '.jpg', '.jpeg'}:
            ocr_text = OCRService._ocr_standalone_image(file_path, languages)
            if ocr_text:
                return ocr_text
            return (
                'Image received but OCR text extraction is unavailable. '
                'Install and configure Tesseract with Hindi language data.'
            )

        return 'Unsupported file type for OCR extraction.'
