import io
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from ..config import get_settings
from ..schemas import DocumentResult

settings = get_settings()


def _resolve_unicode_pdf_font() -> str:
    preferred_path = (settings.pdf_font_path or '').strip()

    candidate_paths = [
        preferred_path,
        '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
        '/Library/Fonts/Arial Unicode.ttf',
    ]

    font_name = 'DafterDocsUnicode'
    for raw_path in candidate_paths:
        if not raw_path:
            continue
        path = Path(raw_path)
        if not path.exists():
            continue
        try:
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name
        except Exception:
            continue

    return 'Helvetica'


def build_output_text(result: DocumentResult) -> str:
    lines = [
        'DAFTERDOCS DOCUMENT OUTPUT',
        '',
        f'Document ID: {result.id}',
        f'Source Language: {result.sourceLanguage}',
        f'Target Language: {result.targetLanguage}',
        f'Generated At: {result.createdAt.isoformat()}',
        '',
        '--- EXTRACTED TEXT ---',
        result.originalText,
        '',
        '--- TRANSLATED TEXT ---',
        result.translatedText,
        '',
        '--- SUMMARY ---',
    ]
    lines.extend([f'{index + 1}. {point}' for index, point in enumerate(result.summary)])
    lines.extend(['', '--- CONCLUSION ---', result.conclusion, ''])
    return '\n'.join(lines)


def build_txt_bytes(result: DocumentResult) -> bytes:
    return build_output_text(result).encode('utf-8')


def build_docx_bytes(result: DocumentResult) -> bytes:
    document = DocxDocument()
    document.add_heading('DafterDocs Document Output', level=1)
    document.add_paragraph(f'Document ID: {result.id}')
    document.add_paragraph(f'Source Language: {result.sourceLanguage}')
    document.add_paragraph(f'Target Language: {result.targetLanguage}')
    document.add_paragraph('')

    document.add_heading('Extracted Text', level=2)
    document.add_paragraph(result.originalText)

    document.add_heading('Translated Text', level=2)
    document.add_paragraph(result.translatedText)

    document.add_heading('Summary', level=2)
    for point in result.summary:
        document.add_paragraph(point, style='List Bullet')

    document.add_heading('Conclusion', level=2)
    document.add_paragraph(result.conclusion)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_bytes(result: DocumentResult) -> bytes:
    text_lines = build_output_text(result).splitlines() or ['']

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    font_name = _resolve_unicode_pdf_font()
    font_size = 11

    width, height = A4
    x = 40
    y = height - 40
    max_width = width - (2 * x)
    line_height = 15
    bottom_margin = 40

    pdf.setFont(font_name, font_size)

    for line in text_lines:
        wrapped_lines = simpleSplit(line or ' ', font_name, font_size, max_width) or [' ']
        for wrapped in wrapped_lines:
            if y < bottom_margin:
                pdf.showPage()
                pdf.setFont(font_name, font_size)
                y = height - 40
            pdf.drawString(x, y, wrapped)
            y -= line_height

        if line == '':
            y -= 4
        if y < bottom_margin:
            pdf.showPage()
            pdf.setFont(font_name, font_size)
            y = height - 40

    pdf.save()
    return buffer.getvalue()
